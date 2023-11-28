
"""
general doc:
this file holds the conenction handler and general system functions such as saving files, and geenrating reports:
for updating devices and rolling-back in batch, which is an extended feature, see the next doc string below
"""
"""
for rollback functionality: According to author of NAPALM libarary this is how it's done:
(https://pynet.twb-tech.com/blog/adding-cisco-ios-to-napalm.html)
"using candidate configs stored in the routers The rollback command will cause 
the 'rollback_config.txt' file to become the running-config (once again using Cisco's 'configure replace' command (think in IOS-XE terms)). 
Note, the current running-config is saved to rollback_config.txt on any commit operation."

NAPALM is a major automation framework for routers and switches. It is reliable.

Also note I'm working on adding  a multiconnect dynamic screen as web hosted for batch updating configs - right now CLI supported,
however until I add something better (integrated Putty, integrated multi-CLi via web page, then 
for now this program is great for obtaining routing and device diagnostics saving it to (MongoDB?)
and comparing results after config changes, notably via traditional methods (the device CLI or programatically)
"""

from napalm import get_network_driver
from netmiko import ConnectHandler
from netmiko import NetMikoTimeoutException, NetMikoAuthenticationException
import threading
import time
import ipaddress
import time
import datetime
import os
import json
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

#dictionaries  -> the only needed dictionary is device_conenctions the rest are defined in the view layer
device_connections = {}
# I'm including command dicts now for when we expand the program to other than cisco devices
repeat = 3
timeout = 2
pings_dict = {'ios': 'ping {} repeat {} timeout {}', 'junos': 'ping {} count {} wait {}', }
bgp_summary_dict = {"ios": "show ip bgp", "junos": "show route protocol bgp", "eos" : "show ip bgp"}


#this reads the infile
def device_list_populator(infile = "device_list.txt"):
    device_list = open(infile, "r")
    device_list = [line.strip('[]\n').split(", ") for line in device_list] # Split it into variables
    device_list = [[element.strip('"') for element in line] for line in device_list]
    return device_list


#handle_ip_addresses() is basically testing pings: it is threaded under appending_to_ping_table()
def handle_ip_addresses(device, ip_addresses, driver_type, output_dictionary={}):
    for ip_address in ip_addresses:
        ip_network = ipaddress.ip_network(ip_address.strip())
        for ip in ip_network.hosts():
            ip_str = str(ip)
            try:
                ping_output = device._send_command(
                    pings_dict[driver_type].format(ip_str, repeat, timeout)
                )
                if "!" in ping_output or "from" in ping_output: #"from" is for junos or Linux: I may change this to a "sucess" on system call
                    output_dictionary.setdefault(device.hostname, []).append(ip_str)
                    print(f"Thread-{threading.current_thread().name}: {output_dictionary}") #this helps user see that his threads are running
                    #threading.enumerate() -> this is another way to do the above
                    #threading.active_count()
            except (NetMikoTimeoutException, NetMikoAuthenticationException) as e:
                print(f"Error while pinging {ip_str} on {device.hostname}: {str(e)}")

#for appending_ping_results in the main layer of our program, we feed in the optional output_dictionary to store pings,
#this is the diagnostic that takes longest, however it is threaded (run in parallel while considering python interpreter lock)
def appending_to_ping_table(ip_address_list, device_connections, output_dictionary={}):
    ping_thread_list = []
    for device, (driver_type, device_ip, username, password, secret) in device_connections.items():
        pinging_thread = threading.Thread(target=handle_ip_addresses, args=(device, ip_address_list, driver_type, output_dictionary))
        ping_thread_list.append(pinging_thread)
        pinging_thread.start()
    for ping in ping_thread_list:
        print(ping.name) #Printing and verifying the threads are working
        ping.join()

#this function populates our device_conenctions list, with conenction variables (username, pass, ip, device_type)
# all of our other functions use these variables
def initialize_device_connections(item,):
    #we unpack the .txt file of devices to get the variables below
    device_type, device_ip, username, password, secret = item
    driver = get_network_driver(device_type) #NAPALMs built-in get_network_driver

    print("The driver: " +  str(driver))
    napalm_device = driver(device_ip, username, password, optional_args = {"read_timeout":50,"global_delay_factor":8,"dest_file_system":"disk0:","auto_file_prompt":True,"secret":secret,"inline_transfer":True})
    #the dest_file_system is where to save rollback and merge. For Cisco devices make sure to enable SCP, file transfers in netmiko are done via SCP.
    #the auto_file_prompt option is to override prompts for file transfers on the device: typically not needed unless set explicitly on the remote device,
    #the auto_file_prompt is only used for batch updating devices, which this program can do, but not overall intent
    #delay_factor to avoid premature errors.

    napalm_device.open() #NAPALMs open method: This is the point in our program where we enable the device for the first time
    # napalm_device.device._send_command("enable") # optional send commands
    # we update a dictionary with the napalm device object, and the connection variables
    device_connections[napalm_device] = device_type, device_ip, username, password, secret


#this threads threaded_network_driver and is an API to initialize_device_connections
def threaded_network_driver(device_list,):
    threads = []
    for item in device_list:
        thread = threading.Thread(target=initialize_device_connections, args=(item,))
        threads.append(thread)
        thread.start()
        time.sleep(.5) #simulates concurrency: introducing a for loop below is what does this, however this time.sleep helps if threaded_network_driver is called immediately after threading the device connections
    for thread in threads:
        thread.join()
        # print(device_connections)

#this function we pass in same number of dictionaries and same number of filenames as lists.
# If filenames is empty the file is result_timestamp.txt for result in results_list
def save_results_to_file(results_list=None, filenames=None):
    output_directory = 'diagnostics_results'
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M")

    if not os.path.exists(output_directory):
        os.makedirs(output_directory)

    if not isinstance(results_list, list):
        results_list = [results_list]

    if not filenames:
        filenames = [None] * len(results_list)
    elif not isinstance(filenames, list):
        filenames = [filenames]

    if len(filenames) != len(results_list):
        print("Error: Number of dictionaries and filenames must be the same.")
        return

    for results, filename in zip(results_list, filenames):
        if not filename:
            filename = results.__name__ if callable(results) else results.__class__.__name__
        file_path = os.path.join(output_directory, f"{filename}_{timestamp}.txt")

        if callable(results):
            # If the result is a function, call it to get the actual dictionary
            results = results()

        with open(file_path, 'w') as f:
            json.dump(results, f, indent=4)  # Save the dictionary as a JSON-formatted string

        print(f"Results saved in {file_path}")
"""
We only use the below if conencting to more than 200-300+ devices:
max_workers = 100
semaphore = threading.Semaphore(max_workers)
place this under threaded_ping_results()  --> semaphore.release() 
"""


"""
    This section of below code involves report generation via word docs: we can modify this to PDF
    I prefer word docs
"""





def generate_report_difference(routing_differences={}, ping_differences={}, napalm_differences={}):
    # Create a Word document
    document = Document()

    # Group differences by IP address
    all_differences = {"Routing": routing_differences, "Ping": ping_differences, "Napalm": napalm_differences}

    for source, differences in all_differences.items():
        document.add_heading(f"{source} Differences", level=1)

        # Group differences by IP address
        grouped_differences = {}
        for difference in differences:
            ip_address = difference['ip_address']
            if ip_address not in grouped_differences:
                grouped_differences[ip_address] = []
            grouped_differences[ip_address].append(difference)

        # Print differences for each IP address
        for ip_address, ip_differences in grouped_differences.items():
            document.add_heading(f"IP Address: {ip_address}", level=2)
            for ip_difference in ip_differences:
                document.add_paragraph(str(ip_difference))

    # Save the document
    current_directory = os.getcwd()
    report_path = os.path.join(current_directory, 'diagnostics_results', 'comparison_report.docx')
    document.save(report_path)


#big Note: We can merge the two report generation functions: Splitting them up gives us a little bit more manual control:
#also I wish to move the directory =, and ping-file, napalm_file, routing_file procedures outside of this function at some point
#also we could pass in dicts directly - is smarter to generate based on saved files

def generate_report_by_ip_address():
    # Specify the directory where the files are located
    directory = 'diagnostics_results'

    # List all files in the directory
    all_files = os.listdir(directory)

    # Filter files based on their names
    ping_files = [file for file in all_files if file.startswith('ping_results_before') and file.endswith('.txt')]
    napalm_files = [file for file in all_files if file.startswith('napalm_results_dict') and file.endswith('.txt')]
    routing_files = [file for file in all_files if file.startswith('routing_dict_before') and file.endswith('.txt')]

    # Get the latest file for each type
    latest_ping_file = max(ping_files) if ping_files else None
    latest_napalm_file = max(napalm_files) if napalm_files else None
    latest_routing_file = max(routing_files) if routing_files else None

    # Check if at least one file of each type exists
    if latest_ping_file and latest_napalm_file and latest_routing_file:
        # Construct full file paths
        ping_results_file = os.path.join(directory, latest_ping_file)
        napalm_results_file = os.path.join(directory, latest_napalm_file)
        routing_results_file = os.path.join(directory, latest_routing_file)

        # Read files and generate report
        with open(ping_results_file, 'r') as f:
            ping_results = json.load(f)

        with open(napalm_results_file, 'r') as f:
            napalm_results = json.load(f)

        with open(routing_results_file, 'r') as f:
            routing_results = json.load(f)

        # Create a Word document
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M")
        document = Document()

        # Group results by IP address
        all_results = {"Ping": ping_results, "Napalm": napalm_results, "Routing": routing_results}

        for source, results in all_results.items():
            document.add_heading(f"{source} Results Grouped by IP Address", level=1)

            # Group results by IP address
            grouped_results = {}
            no_default_route_ips = []  # List to store IP addresses without 0.0.0.0 default route
            for ip_address, result_data in results.items():
                if ip_address not in grouped_results:
                    grouped_results[ip_address] = []
                grouped_results[ip_address].append(result_data)

                # Check if IP address has no 0.0.0.0 default route
                has_default_route = any('ip_address' in entry and entry['ip_address'] == '0.0.0.0/0' for entry in result_data)
                if not has_default_route:
                    no_default_route_ips.append(ip_address)

            # Print results for each IP address
            for ip_address, ip_results in grouped_results.items():
                document.add_heading(f"IP Address: {ip_address}", level=2)
                combined_results = []  # List to store combined results for each IP address

                for ip_result in ip_results:
                    if isinstance(ip_result, list):
                        # If ip_result is a list, concatenate each item in the list
                        combined_results.extend(ip_result)
                    elif isinstance(ip_result, dict):
                        # If ip_result is a dictionary, add a formatted paragraph for each key-value pair
                        for key, value in ip_result.items():
                            if source == "Napalm":
                                # Bold the NAPALM function names
                                paragraph = f"{key}: {json.dumps(value, indent=2)}"
                                combined_results.append(paragraph)

                # Add the combined results as a single paragraph
                document.add_paragraph('\n'.join(combined_results))

            # Highlight or bold IP addresses without 0.0.0.0 default route
            if no_default_route_ips:
                document.add_heading("IP Addresses Without 0.0.0.0 Default Route", level=2).bold = True
                for ip in no_default_route_ips:
                    document.add_paragraph(f"{ip} does not have associated 0.0.0.0 default-route").bold = True

        # Save the document
        current_directory = os.getcwd()
        report_path = os.path.join(current_directory, 'diagnostics_results', f'results_report_{timestamp}.docx')
        document.save(report_path)

# Example usage:

